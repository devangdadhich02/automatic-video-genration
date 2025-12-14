from __future__ import annotations

import os
from datetime import datetime, timezone
from pathlib import Path
from typing import Iterable, Literal, Optional

from .config import settings

ExportFormat = Literal["txt", "docx", "pdf"]


def _utc_now_compact() -> str:
    return datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")


def _ensure_exports_dir() -> Path:
    p = Path(settings.EXPORTS_DIR)
    p.mkdir(parents=True, exist_ok=True)
    return p


def export_script(
    *,
    content_id: str,
    title: str,
    text: str,
    formats: Iterable[ExportFormat],
    filename_prefix: Optional[str] = None,
) -> dict:
    """Export generated script to one or more formats (best-effort).

    Always produces TXT (even if other formats fail) to satisfy "no error" expectation.
    """

    exports_dir = _ensure_exports_dir()
    safe_prefix = (filename_prefix or "script").strip().replace(" ", "_")
    stamp = _utc_now_compact()
    base = f"{safe_prefix}_{content_id}_{stamp}"

    out: dict = {"content_id": content_id, "exports": [], "errors": []}

    wanted = list(dict.fromkeys([f for f in formats if f]))  # de-dupe, preserve order
    if not wanted:
        wanted = ["txt"]

    # Always write TXT as fallback artifact
    try:
        txt_path = exports_dir / f"{base}.txt"
        txt_path.write_text(text or "", encoding="utf-8")
        out["exports"].append({"format": "txt", "path": str(txt_path)})
    except Exception as e:
        out["errors"].append(f"txt_export_failed: {e}")

    if "docx" in wanted:
        try:
            from docx import Document  # type: ignore

            doc = Document()
            if title:
                doc.add_heading(title, level=1)
            for para in (text or "").split("\n\n"):
                p = (para or "").strip()
                if p:
                    doc.add_paragraph(p)
            docx_path = exports_dir / f"{base}.docx"
            doc.save(str(docx_path))
            out["exports"].append({"format": "docx", "path": str(docx_path)})
        except Exception as e:
            out["errors"].append(f"docx_export_failed: {e}")

    if "pdf" in wanted:
        try:
            # fpdf2 is lightweight and works offline.
            from fpdf import FPDF  # type: ignore

            pdf = FPDF()
            pdf.set_auto_page_break(auto=True, margin=15)
            pdf.add_page()
            pdf.set_font("Helvetica", size=12)

            if title:
                pdf.set_font("Helvetica", style="B", size=14)
                pdf.multi_cell(0, 8, title)
                pdf.ln(2)
                pdf.set_font("Helvetica", size=12)

            # Use multi_cell to wrap text
            for line in (text or "").splitlines():
                pdf.multi_cell(0, 6, line)
            pdf_path = exports_dir / f"{base}.pdf"
            pdf.output(str(pdf_path))
            out["exports"].append({"format": "pdf", "path": str(pdf_path)})
        except Exception as e:
            out["errors"].append(f"pdf_export_failed: {e}")

    return out


